import pytesseract
from pdf2image import convert_from_path
from docx import Document

# Optional: Set tesseract path if it's not in your system PATH
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

pdf_path = r"C:\Users\saqom\Downloads\Telegram Desktop\zangak7.pdf"
output_docx_path = r"C:\Users\saqom\Downloads\Telegram Desktop\zangak7.docx"

# Convert PDF pages to images using poppler_path
images = convert_from_path(pdf_path, dpi=500, poppler_path=r"C:\poppler\Library\bin")

# Create a new Word document
doc = Document()
doc.add_heading('OCR Extracted Text', level=1)

# OCR each image and add text to the Word file
for i, image in enumerate(images):
    text = pytesseract.image_to_string(image, lang='hye')  # 'hye' is Armenian
    doc.add_paragraph(f'--- Page {i + 1} ---\n{text}')

# Save the document
doc.save(output_docx_path)

print("OCR completed and saved to Word document.")